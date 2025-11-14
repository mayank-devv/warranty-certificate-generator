import io
import re
from datetime import datetime
import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ----------------------------------------------------
# Streamlit UI
# ----------------------------------------------------
st.set_page_config(page_title="Warranty Certificate Generator", page_icon="🧾", layout="centered")
st.title("🧾 Warranty Certificate Generator – PDF Auto Extract")
st.caption("Upload GEMC PDF + DOCX template → Auto-formatted Warranty Certificate")

# ----------------------------------------------------
# File Inputs
# ----------------------------------------------------
template_file = st.file_uploader("Upload Warranty DOCX Template", type=["docx"])
pdf_file = st.file_uploader("Upload GEMC PDF (Auto Extract)", type=["pdf"])

submitted = st.button("Generate Certificate")


# ----------------------------------------------------
# Utility Styling Functions
# ----------------------------------------------------
BLUE = RGBColor(0, 112, 192)

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="6" w:space="1" w:color="0070C0"/></w:pBdr>' % nsdecls("w"))
    pPr.append(pBdr)

def render_labeled_paragraph(p, text):
    for i in range(len(p.runs)-1, -1, -1):
        p._element.remove(p.runs[i]._element)

    if ":" in text:
        label, _, value = text.partition(":")
        r1 = p.add_run(label.strip() + ":")
        r1.font.bold = True
        r1.font.color.rgb = BLUE
        r1.font.size = Pt(12)
        r1.font.name = "Calibri"

        r2 = p.add_run(" " + value.strip())
        r2.font.bold = False
        r2.font.color.rgb = BLUE
        r2.font.size = Pt(12)
        r2.font.name = "Calibri"

    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ----------------------------------------------------
# Text Extraction Patterns
# ----------------------------------------------------
def extract_key(pattern, text):
    match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
    return match.group(1).strip() if match else ""


def extract_from_pdf(pdf):
    """Reads PDF & returns raw text"""
    full_text = ""
    with pdfplumber.open(pdf) as pdf_doc:
        for page in pdf_doc.pages:
            full_text += page.extract_text() + "\n"
    return full_text


def auto_extract_fields(raw):
    """Extracts all GEMC fields automatically"""

    data = {}

    # 🔹 Basic blocks
    data["GEMContractNo"] = extract_key(r"(GEMC-[0-9\-]+)", raw)
    data["Date"] = extract_key(r"Generated Date\s*:\s*([0-9A-Za-z\-]+)", raw)

    # 🔹 Product Area
    data["Brand"] = extract_key(r"Brand\s*:\s*(.+)", raw)
    data["Model"] = extract_key(r"Model\s*:\s*([A-Za-z0-9\-\s]+)", raw)
    data["ProductName"] = extract_key(r"Product Name\s*:\s*(.+)", raw)

    # 🔹 Category
    data["Category"] = extract_key(r"Category Name.*?:\s*(.+)", raw)
    if not data["Category"]:
        if "AC" in data["ProductName"].upper():
            data["Category"] = "AC"

    # 🔹 Quantity
    data["Quantity"] = extract_key(r"(\d+)\s*pieces", raw)
    if data["Quantity"]:
        data["Quantity"] += " Unit"
    else:
        data["Quantity"] = "1 Unit"

    # 🔹 Warranty
    data["Warranty"] = extract_key(r"Comprehensive Warranty.*?\(?in years\)?\s*([0-9]+)", raw)
    if data["Warranty"]:
        data["Warranty"] += " Years Overall"

    data["WarrantyOnCompressor"] = extract_key(
        r"Warranty on Compressor.*?\(?in years\)?\s*([0-9]+)", raw)

    # 🔹 Customer / Organisation / Address
    data["CustomerName"] = extract_key(r"Designation\s*:\s*(.+)", raw)
    data["Organisation"] = extract_key(r"Organisation Name\s*:\s*(.+)", raw)

    # Address - multiple lines
    addr_block = extract_key(r"Address\s*:\s*(.+?)(?=Email|GSTIN|$)", raw)
    addr_block = addr_block.replace("\n", ", ").replace(",,", ",")
    data["Address"] = addr_block.strip(", ").strip()

    # 🔹 Serial Number Detection (optional)
    serials = re.findall(r"[A-Z0-9]{12,25}", raw)
    data["SerialNumber"] = serials[0] if serials else ""

    return data


# ----------------------------------------------------
# Merge & Replace Template
# ----------------------------------------------------
def merge_and_replace(doc, mapping):

    def _process(container):
        for p in container.paragraphs:
            original = "".join(run.text for run in p.runs)
            replaced = original
            for k, v in mapping.items():
                replaced = replaced.replace(k, v)
            render_labeled_paragraph(p, replaced)

        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    _process(cell)

    _process(doc)


# ----------------------------------------------------
# MAIN PROCESS
# ----------------------------------------------------
if submitted:
    if not template_file:
        st.error("Upload DOCX template first.")
    elif not pdf_file:
        st.error("Upload GEMC PDF.")
    else:
        # 🔹 Extract PDF text
        raw_text = extract_from_pdf(pdf_file)

        # 🔹 Extract fields
        data = auto_extract_fields(raw_text)

        # 🔹 Load template
        doc = Document(template_file)

        # 🔹 Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # 🔹 Mapping
        mapping = {
            "{Company}": "Shrii Salez Corporation",
            "{Category}": data["Category"],
            "{Brand}": data["Brand"],
            "{Make}": data["Brand"],
            "{ProductName}": data["ProductName"],
            "{Model}": data["Model"],
            "{Quantity}": data["Quantity"],
            "{SerialNumber}": data["SerialNumber"],
            "{GEMContractNo}": data["GEMContractNo"],
            "{Warranty}": data["Warranty"],
            "{WarrantyOnCompressor}": data["WarrantyOnCompressor"],
            "{CustomerName}": data["CustomerName"],
            "{Organisation}": data["Organisation"],
            "{Address}": data["Address"],
            "{Date}": data["Date"],
        }

        # 🔹 Replace in DOCX
        merge_and_replace(doc, mapping)

        # 🔹 Enhance formatting (same as previous)
        if doc.paragraphs:
            header = doc.paragraphs[0]
            for run in header.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = BLUE
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Center letterhead lines
        for i in range(1, 7):
            if i < len(doc.paragraphs):
                p = doc.paragraphs[i]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.color.rgb = BLUE
                    run.font.size = Pt(12)
                    run.font.name = "Calibri"

        # Add blue line below address block
        for i, p in enumerate(doc.paragraphs):
            if "@" in p.text or "Email" in p.text:
                new_p = doc.paragraphs[i + 1].insert_paragraph_before("")
                add_horizontal_line(new_p)
                break

        # Format WARRANTY CERTIFICATE heading
        for p in doc.paragraphs:
            if "WARRANTY CERTIFICATE" in p.text.upper():
                for run in p.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.underline = True
                    run.font.color.rgb = BLUE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save output
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)

        output_name = f"Warranty_{data['CustomerName'].replace(' ','_')}_{data['GEMContractNo']}.docx"

        st.success("✅ Certificate Generated Successfully")
        st.download_button("⬇ Download Warranty Certificate", data=out,
                          file_name=output_name,
                          mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
